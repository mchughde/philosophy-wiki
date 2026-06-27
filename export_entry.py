"""
export_entry.py — Philosophy Wiki → 6x9 Word document
Usage: python3 export_entry.py <slug>
Example: python3 export_entry.py socrates
"""

import sys, os, re, json, glob, zipfile

# ── locate wiki folder (script lives in the wiki root) ────────────────────────
WIKI_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE_DIR = os.path.join(WIKI_DIR, "Source material")
ENTRIES_JS = os.path.join(WIKI_DIR, "entries.js")

# ── parse entries.js ──────────────────────────────────────────────────────────

def load_entries():
    with open(ENTRIES_JS, encoding="utf-8") as f:
        raw = f.read()
    # Strip JS wrapper: remove 'const ENTRIES = ' and 'export default ENTRIES;'
    raw = re.sub(r"^\s*//.*$", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"const ENTRIES\s*=\s*", "", raw)
    raw = re.sub(r"export default ENTRIES\s*;?", "", raw)
    raw = raw.strip().rstrip(";").strip()
    return json.loads(raw)

# ── markdown helpers ──────────────────────────────────────────────────────────

def parse_inline(text):
    """
    Return a list of (run_text, is_italic) tuples from a markdown string.
    Handles *italic* and [[wikilinks]] (rendered as plain text).
    """
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)   # strip [[links]]
    parts = []
    pattern = re.compile(r"\*([^*]+)\*")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts if parts else [(text, False)]

def parse_body(body):
    """
    Return (intro_paragraphs, sections) where:
      intro_paragraphs = [str, ...]
      sections = [{"title": str, "paragraphs": [str, ...]}, ...]
    Skips any section titled 'Quotes'.
    """
    # Normalise line endings
    body = body.replace("\\n", "\n")

    # Split on ## headings
    parts = re.split(r"\n##\s+", body)
    intro_block = parts[0].strip()
    intro_paragraphs = [p.strip() for p in intro_block.split("\n\n") if p.strip()]

    sections = []
    for part in parts[1:]:
        lines = part.split("\n", 1)
        title = lines[0].strip()
        if title.lower() == "quotes":
            continue
        content = lines[1].strip() if len(lines) > 1 else ""
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        sections.append({"title": title, "paragraphs": paragraphs})

    return intro_paragraphs, sections

# ── image lookup ──────────────────────────────────────────────────────────────

def find_local_image(title):
    """
    Look for an image file in any Source material subfolder whose name contains
    the philosopher's title (case-insensitive).
    Returns the path or None.
    """
    IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    pattern = os.path.join(SOURCE_DIR, "*")
    for folder in glob.glob(pattern):
        if not os.path.isdir(folder):
            continue
        if title.lower() not in os.path.basename(folder).lower():
            continue
        for fname in os.listdir(folder):
            if os.path.splitext(fname)[1].lower() in IMAGE_EXTS:
                return os.path.join(folder, fname)
    return None

def download_image(url, dest_path):
    """Download an image from a URL to dest_path. Returns True on success."""
    try:
        import urllib.request
        headers = {"User-Agent": "Mozilla/5.0"}
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=15) as resp:
            with open(dest_path, "wb") as f:
                f.write(resp.read())
        return True
    except Exception as e:
        print(f"  Warning: could not download image ({e})")
        return False

# ── date formatting ───────────────────────────────────────────────────────────

def fmt_date(year):
    if year is None:
        return "?"
    if year < 0:
        return f"{abs(year)} BC"
    return f"{year} CE"

def fmt_date_range(born, died):
    return f"{fmt_date(born)} – {fmt_date(died)}"

# ── Word document builder ─────────────────────────────────────────────────────

def build_docx(entry, image_path, out_path):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    title = entry["title"]
    date_range = fmt_date_range(entry.get("born"), entry.get("died"))
    tags_str = "  ·  ".join(t.upper() for t in entry.get("tags", []))
    intro_paras, sections = parse_body(entry["body"])

    # ── helpers ──

    def add_page_number(paragraph):
        run = paragraph.add_run()
        for tag, text in [("w:fldChar", None), ("w:instrText", "PAGE"), ("w:fldChar", None)]:
            el = OxmlElement(tag)
            if tag == "w:instrText":
                el.text = text
            elif tag == "w:fldChar":
                el.set(qn("w:fldCharType"), "begin" if not hasattr(add_page_number, "_end") else "end")
                if not hasattr(add_page_number, "_end"):
                    add_page_number._end = True
                else:
                    delattr(add_page_number, "_end")
            run._r.append(el)
        run.font.name = "Georgia"
        run.font.size = Pt(9)

    def add_page_num(paragraph):
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = "Georgia"
        run.font.size = Pt(9)

    AFTER_PBDR = {"shd","tabs","suppressAutoHyphens","kinsoku","wordWrap",
                  "overflowPunct","topLinePunct","autoSpaceDE","autoSpaceDN",
                  "bidi","adjustRightInd","snapToGrid","spacing","ind",
                  "contextualSpacing","mirrorIndents","suppressOverlap","jc",
                  "textDirection","textAlignment","textboxTightWrap","outlineLvl",
                  "divId","cnfStyle","rPr"}

    def add_border(paragraph, side, color="AAAAAA", sz="4", space="4"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), color)
        pBdr.append(el)
        insert_before = next(
            (c for c in pPr if (c.tag.split("}")[1] if "}" in c.tag else c.tag) in AFTER_PBDR),
            None
        )
        if insert_before is not None:
            pPr.insert(list(pPr).index(insert_before), pBdr)
        else:
            pPr.append(pBdr)

    def body_para(doc, text, first_indent=False, space_before=0, space_after=8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if first_indent:
            p.paragraph_format.first_line_indent = Inches(0.2)
        for fragment, italic in parse_inline(text):
            r = p.add_run(fragment)
            r.font.name   = "Georgia"
            r.font.size   = Pt(10.5)
            r.font.italic = italic
        return p

    def section_hdr(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(text)
        r.font.name      = "Georgia"
        r.font.size      = Pt(7.5)
        r.font.bold      = True
        r.font.all_caps  = True
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        return p

    # ── build ──

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(6)
    sec.page_height   = Inches(9)
    sec.top_margin    = Inches(0.875)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.0)    # inside (spine) margin — wider for binding
    sec.right_margin  = Inches(0.625)  # outside margin

    # Header
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(title.upper())
    hr.font.name      = "Georgia"
    hr.font.size      = Pt(7.5)
    hr.font.all_caps  = True
    hr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    add_border(hp, "bottom", color="CCCCCC", sz="4", space="4")

    # Footer
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_num(fp)

    # Title block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    r.font.name = "Georgia"
    r.font.size = Pt(26)
    r.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(date_range)
    r.font.name      = "Georgia"
    r.font.size      = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    r = p.add_run(tags_str)
    r.font.name      = "Georgia"
    r.font.size      = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Image (if available)
    if image_path:
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ip.paragraph_format.space_before = Pt(6)
        ip.paragraph_format.space_after  = Pt(10)
        ip.add_run().add_picture(image_path, width=Inches(1.75), height=Inches(2.33))

    # Introduction (first paragraph no indent, rest indented)
    for i, para in enumerate(intro_paras):
        body_para(doc, para, first_indent=(i > 0), space_after=10)

    # Sections
    for section in sections:
        section_hdr(doc, section["title"])
        for i, para in enumerate(section["paragraphs"]):
            body_para(doc, para, first_indent=(i > 0))

    # Save and patch zoom bug
    doc.save(out_path)
    _patch_zoom(out_path)
    print(f"  Saved: {out_path}")

def merge_docx_files(docx_paths, out_path):
    """
    Merge a list of single-entry docx files into one combined document.
    Each philosopher gets their own section (starting on an odd/right-hand page)
    with their own header and correctly remapped images.
    Page numbers run continuously throughout.
    """
    from lxml import etree
    import copy

    W           = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R           = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    HEADER_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
    FOOTER_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
    IMAGE_TYPE  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
    CT_NS       = "http://schemas.openxmlformats.org/package/2006/content-types"
    HEADER_CT   = "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"

    # Load all source zips into memory
    sources = []
    for path in docx_paths:
        with zipfile.ZipFile(path, "r") as z:
            sources.append({name: z.read(name) for name in z.namelist()})

    base = sources[0]
    base_doc  = etree.fromstring(base["word/document.xml"])
    base_rels = etree.fromstring(base["word/_rels/document.xml.rels"])
    base_ct   = etree.fromstring(base["[Content_Types].xml"])
    base_body = base_doc.find(f"{{{W}}}body")

    # Pick a safe starting rId above all existing ones in the base
    next_id = max(
        (int(r.get("Id", "rId0").replace("rId", "")) for r in base_rels if r.get("Id","").startswith("rId")),
        default=10
    ) + 1
    # Also track a global image counter so every media file gets a unique name
    img_counter = [1]
    # Seed with any media files already in the base
    for name in base:
        if name.startswith("word/media/"):
            img_counter[0] += 1

    # Strip all existing header refs and header content-type overrides from base
    for rel in [r for r in list(base_rels) if r.get("Type") == HEADER_TYPE]:
        base_rels.remove(rel)
    for el in [e for e in list(base_ct) if "header" in e.get("PartName","").lower()]:
        base_ct.remove(el)

    # Collect extensions already registered as Default content types
    registered_exts = {
        el.get("Extension","").lower()
        for el in base_ct
        if el.tag == f"{{{CT_NS}}}Default"
    }

    # Clear base body — we'll rebuild it from scratch
    for child in list(base_body):
        base_body.remove(child)

    new_files = {}  # extra parts to add to the combined zip

    for i, src in enumerate(sources):
        is_last = (i == len(sources) - 1)

        # Parse this source's relationships
        src_rels = etree.fromstring(src["word/_rels/document.xml.rels"])

        # ── remap image relationships ──────────────────────────────────────────
        # Build old_rId → new_rId map for every image in this source
        rid_map = {}
        for rel in src_rels:
            if rel.get("Type") != IMAGE_TYPE:
                continue
            old_rid    = rel.get("Id")
            old_target = rel.get("Target")                # e.g. "media/image1.jpg"
            ext        = old_target.rsplit(".", 1)[-1].lower()
            new_target = f"media/image{img_counter[0]}.{ext}"
            img_counter[0] += 1
            new_rid    = f"rId{next_id}"; next_id += 1
            rid_map[old_rid] = new_rid

            # Copy image bytes under the new name
            src_img_path = f"word/{old_target}"
            if src_img_path in src:
                new_files[f"word/{new_target}"] = src[src_img_path]

            # Register relationship in combined doc
            etree.SubElement(base_rels, "Relationship",
                Id=new_rid, Type=IMAGE_TYPE, Target=new_target)

            # Register content type if not already present
            if ext not in registered_exts:
                mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
                        "png": "image/png", "gif": "image/gif",
                        "webp": "image/webp"}.get(ext, f"image/{ext}")
                etree.SubElement(base_ct, f"{{{CT_NS}}}Default",
                    Extension=ext, ContentType=mime)
                registered_exts.add(ext)

        # ── extract + patch body elements ──────────────────────────────────────
        src_doc  = etree.fromstring(src["word/document.xml"])
        src_body = src_doc.find(f"{{{W}}}body")
        children = list(src_body)
        sectpr   = copy.deepcopy(children[-1])
        body_els = [copy.deepcopy(e) for e in children[:-1]]

        # Patch r:embed / r:id attributes in body elements to use new rIds
        if rid_map:
            for el in body_els:
                for node in el.iter():
                    for attr in (f"{{{R}}}embed", f"{{{R}}}id", f"{{{R}}}link"):
                        val = node.get(attr)
                        if val and val in rid_map:
                            node.set(attr, rid_map[val])

        # ── header ────────────────────────────────────────────────────────────
        src_hdr_file = None
        for rel in src_rels:
            if rel.get("Type") == HEADER_TYPE:
                src_hdr_file = "word/" + rel.get("Target")
                break

        new_hdr_rid = f"rId{next_id}"; next_id += 1
        new_hdr     = f"header{i + 1}.xml"
        if src_hdr_file and src_hdr_file in src:
            new_files[f"word/{new_hdr}"] = src[src_hdr_file]

        etree.SubElement(base_rels, "Relationship",
            Id=new_hdr_rid, Type=HEADER_TYPE, Target=new_hdr)
        etree.SubElement(base_ct, f"{{{CT_NS}}}Override",
            PartName=f"/word/{new_hdr}", ContentType=HEADER_CT)

        # ── patch sectPr ──────────────────────────────────────────────────────
        for el in list(sectpr):
            tag = el.tag.split("}")[1] if "}" in el.tag else el.tag
            if tag == "headerReference":
                sectpr.remove(el)
            if tag == "footerReference" and i > 0:
                sectpr.remove(el)

        href = etree.Element(f"{{{W}}}headerReference")
        href.set(f"{{{W}}}type", "default")
        href.set(f"{{{R}}}id", new_hdr_rid)
        sectpr.insert(0, href)

        # ── append to combined body ───────────────────────────────────────────
        for el in body_els:
            base_body.append(el)

        if is_last:
            base_body.append(sectpr)
        else:
            type_el = etree.Element(f"{{{W}}}type")
            type_el.set(f"{{{W}}}val", "oddPage")
            sectpr.insert(0, type_el)
            para = etree.Element(f"{{{W}}}p")
            pPr  = etree.SubElement(para, f"{{{W}}}pPr")
            pPr.append(sectpr)
            base_body.append(para)

    # ── serialise ─────────────────────────────────────────────────────────────
    new_doc_xml  = etree.tostring(base_doc,  xml_declaration=True, encoding="UTF-8", standalone=True)
    new_rels_xml = etree.tostring(base_rels, xml_declaration=True, encoding="UTF-8", standalone=True)
    new_ct_xml   = etree.tostring(base_ct,   xml_declaration=True, encoding="UTF-8", standalone=True)

    # Skip old header files from base (we're replacing them all)
    old_hdrs = {f"word/header{j + 1}.xml" for j in range(len(sources))}

    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in base.items():
            if name in old_hdrs:
                continue
            if name == "word/document.xml":
                zout.writestr(name, new_doc_xml)
            elif name == "word/_rels/document.xml.rels":
                zout.writestr(name, new_rels_xml)
            elif name == "[Content_Types].xml":
                zout.writestr(name, new_ct_xml)
            else:
                zout.writestr(name, data)
        for fname, content in new_files.items():
            zout.writestr(fname, content)

    _patch_zoom(out_path)


def _patch_zoom(path):
    tmp = path + ".tmp"
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                text = data.decode("utf-8")
                text = re.sub(r"(<w:zoom\b)(?![^>]*w:percent)", r'\1 w:percent="100"', text)
                # Enable mirror margins (Word flips inside/outside on alternating pages)
                if "<w:mirrorMargins" not in text:
                    text = text.replace("</w:settings>", "<w:mirrorMargins/></w:settings>")
                data = text.encode("utf-8")
            zout.writestr(item, data)
    os.replace(tmp, path)

# ── main ──────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 export_entry.py <slug>")
        print("Example: python3 export_entry.py socrates")
        sys.exit(1)

    slug = sys.argv[1].lower()

    print(f"Loading entries.js...")
    entries = load_entries()
    entry = next((e for e in entries if e["slug"] == slug), None)
    if not entry:
        available = [e["slug"] for e in entries]
        print(f"No entry found for '{slug}'. Available: {', '.join(available)}")
        sys.exit(1)

    title = entry["title"]
    print(f"Found: {title}")

    # Find image
    image_path = find_local_image(title)
    if image_path:
        print(f"  Image: {os.path.basename(image_path)} (local)")
    elif entry.get("photo"):
        print(f"  Image not found locally — trying URL fallback...")
        tmp_img = os.path.join(WIKI_DIR, f"_tmp_{slug}_img.jpg")
        if download_image(entry["photo"], tmp_img):
            image_path = tmp_img
            print(f"  Image downloaded from URL")
        else:
            print(f"  Continuing without image")
    else:
        print(f"  No image available")

    # Output path
    out_path = os.path.join(WIKI_DIR, f"{slug}_6x9.docx")

    print(f"Building Word document...")
    build_docx(entry, image_path, out_path)

    # Clean up temp image
    tmp_img = os.path.join(WIKI_DIR, f"_tmp_{slug}_img.jpg")
    if os.path.exists(tmp_img):
        os.remove(tmp_img)

    print(f"Done. Open: {out_path}")

if __name__ == "__main__":
    main()
